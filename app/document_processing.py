"""
Document Processing Module for AI Vectorizer

This module provides functionality for text extraction from various document formats,
text cleaning, and chunking for efficient processing and indexing.
"""

import os
import re
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any, Union
import logging
from datetime import datetime

# PDF processing
import fitz  # PyMuPDF

# DOCX processing
import docx

# For tokenization and chunking
import nltk
from nltk.tokenize import sent_tokenize
import tiktoken

# Ensure NLTK data is downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# GPT tokenizer for token counting
ENCODING = tiktoken.get_encoding("cl100k_base")  # Default encoding for newer models

class DocumentProcessingError(Exception):
    """Exception raised for errors in document processing."""
    pass

def extract_text_from_pdf(file_path: Union[str, Path]) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
        
    Raises:
        DocumentProcessingError: If text extraction fails
    """
    try:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
            
        text = ""
        with fitz.open(file_path) as pdf:
            page_count = len(pdf)
            logger.info(f"Processing PDF with {page_count} pages")
            
            for page_num, page in enumerate(pdf):
                text += page.get_text()
                # Add a newline between pages if not already present
                if not text.endswith('\n'):
                    text += '\n'
                    
        if not text.strip():
            logger.warning(f"No text extracted from PDF: {file_path}")
            
        return text
        
    except Exception as e:
        error_msg = f"Error extracting text from PDF {file_path}: {str(e)}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg) from e

def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        DocumentProcessingError: If text extraction fails
    """
    try:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs]
        
        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells if cell.text)
                if row_text:
                    table_texts.append(row_text)
        
        # Combine all text
        all_text = '\n'.join(paragraphs + table_texts)
        
        if not all_text.strip():
            logger.warning(f"No text extracted from DOCX: {file_path}")
            
        return all_text
        
    except Exception as e:
        error_msg = f"Error extracting text from DOCX {file_path}: {str(e)}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg) from e

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw text to clean
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
        
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n+', '\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove non-printable characters
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    return text.strip()

def count_tokens(text: str) -> int:
    """
    Count the number of tokens in a text string using tiktoken.
    
    Args:
        text: Text to count tokens for
        
    Returns:
        Number of tokens
    """
    if not text:
        return 0
        
    tokens = ENCODING.encode(text)
    return len(tokens)

def chunk_text(
    text: str, 
    max_tokens: int = 512, 
    overlap_tokens: int = 50
) -> List[str]:
    """
    Split text into chunks with a specified maximum token count and overlap.
    
    Args:
        text: Text to split into chunks
        max_tokens: Maximum number of tokens per chunk (default: 512)
        overlap_tokens: Number of tokens to overlap between chunks (default: 50)
        
    Returns:
        List of text chunks
    """
    if not text or max_tokens <= 0:
        return []
        
    # Ensure overlap is not larger than max_tokens
    overlap_tokens = min(overlap_tokens, max_tokens // 2)
    
    # Clean the text first
    text = clean_text(text)
    
    # Split text into paragraphs
    paragraphs = [p for p in text.split('\n') if p.strip()]
    
    chunks = []
    current_chunk = ""
    current_tokens = 0
    
    for paragraph in paragraphs:
        # Count tokens in this paragraph
        paragraph_tokens = count_tokens(paragraph)
        
        # If a single paragraph exceeds max_tokens, we need to split it
        if paragraph_tokens > max_tokens:
            # If we have content in the current chunk, add it to chunks
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""
                current_tokens = 0
            
            # Split the paragraph into sentences
            sentences = sent_tokenize(paragraph)
            
            sentence_chunk = ""
            sentence_tokens = 0
            
            for sentence in sentences:
                sentence_token_count = count_tokens(sentence)
                
                # If a single sentence exceeds max_tokens, split it into smaller parts
                if sentence_token_count > max_tokens:
                    # If we have content in the sentence chunk, add it to chunks
                    if sentence_chunk:
                        chunks.append(sentence_chunk)
                        sentence_chunk = ""
                        sentence_tokens = 0
                    
                    # Split the sentence into smaller parts
                    words = sentence.split()
                    word_chunk = ""
                    word_tokens = 0
                    
                    for word in words:
                        word_with_space = word + " "
                        word_token_count = count_tokens(word_with_space)
                        
                        if word_tokens + word_token_count <= max_tokens:
                            word_chunk += word_with_space
                            word_tokens += word_token_count
                        else:
                            chunks.append(word_chunk.strip())
                            word_chunk = word_with_space
                            word_tokens = word_token_count
                    
                    if word_chunk:
                        chunks.append(word_chunk.strip())
                
                # If adding this sentence would exceed max_tokens, start a new chunk
                elif sentence_tokens + sentence_token_count > max_tokens:
                    chunks.append(sentence_chunk.strip())
                    
                    # Start new chunk with overlap from the end of the previous chunk
                    if sentence_tokens > overlap_tokens:
                        # Find the last few sentences that fit within overlap_tokens
                        overlap_text = ""
                        overlap_count = 0
                        for s in reversed(sent_tokenize(sentence_chunk)):
                            s_count = count_tokens(s + " ")
                            if overlap_count + s_count <= overlap_tokens:
                                overlap_text = s + " " + overlap_text
                                overlap_count += s_count
                            else:
                                break
                        
                        sentence_chunk = overlap_text + sentence + " "
                        sentence_tokens = overlap_count + sentence_token_count
                    else:
                        sentence_chunk = sentence + " "
                        sentence_tokens = sentence_token_count
                
                else:
                    sentence_chunk += sentence + " "
                    sentence_tokens += sentence_token_count
            
            if sentence_chunk:
                chunks.append(sentence_chunk.strip())
        
        # If adding this paragraph would exceed max_tokens, start a new chunk
        elif current_tokens + paragraph_tokens > max_tokens:
            chunks.append(current_chunk.strip())
            
            # Start new chunk with overlap from the end of the previous chunk
            if current_tokens > overlap_tokens:
                # Find the last few sentences that fit within overlap_tokens
                overlap_text = ""
                overlap_count = 0
                for s in reversed(sent_tokenize(current_chunk)):
                    s_count = count_tokens(s + " ")
                    if overlap_count + s_count <= overlap_tokens:
                        overlap_text = s + " " + overlap_text
                        overlap_count += s_count
                    else:
                        break
                
                current_chunk = overlap_text + paragraph + "\n"
                current_tokens = overlap_count + paragraph_tokens
            else:
                current_chunk = paragraph + "\n"
                current_tokens = paragraph_tokens
        
        else:
            current_chunk += paragraph + "\n"
            current_tokens += paragraph_tokens
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

def process_document(
    file_path: Union[str, Path],
    max_tokens: int = 512,
    overlap_tokens: int = 50
) -> Dict[str, Any]:
    """
    Process a document: extract text, clean it, and split into chunks.
    
    Args:
        file_path: Path to the document file
        max_tokens: Maximum number of tokens per chunk
        overlap_tokens: Number of tokens to overlap between chunks
        
    Returns:
        Dictionary with processing results including:
        - extracted_text: Full extracted text
        - chunks: List of text chunks
        - metadata: Document metadata (file type, size, etc.)
        
    Raises:
        DocumentProcessingError: If processing fails
    """
    try:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file metadata
        file_stats = file_path.stat()
        file_type = file_path.suffix.lower().lstrip('.')
        
        # Extract text based on file type
        if file_type == 'pdf':
            extracted_text = extract_text_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            extracted_text = extract_text_from_docx(file_path)
        elif file_type in ['txt', 'md', 'py', 'js', 'html', 'css', 'json']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
        else:
            raise DocumentProcessingError(f"Unsupported file type: {file_type}")
        
        # Clean the text
        cleaned_text = clean_text(extracted_text)
        
        # Split into chunks
        chunks = chunk_text(cleaned_text, max_tokens, overlap_tokens)
        
        # Count tokens in the full text
        total_tokens = count_tokens(cleaned_text)
        
        # Prepare result
        result = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_type": file_type,
            "file_size": file_stats.st_size,
            "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "total_tokens": total_tokens,
            "chunk_count": len(chunks),
            "chunks": chunks,
            "extracted_text": cleaned_text,
            "processing_time": datetime.now().isoformat()
        }
        
        return result
        
    except Exception as e:
        error_msg = f"Error processing document {file_path}: {str(e)}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg) from e
